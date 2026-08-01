# -*- coding: utf-8 -*-
"""extract_text.py · 从 PDF/DOCX/TXT 提取纯文本（WF-01 主路径）

用法:
  python tools/extract_text.py --input <file> --output <out.txt>

规则:
  - DOCX: 优先 python-docx；失败降级 zipfile + 正则解析 word/document.xml
  - PDF:  pypdf；扫描件提取为空时明确报错（退出码 2），提示另存为 txt
  - TXT:  utf-8 直读，失败容错 gb18030
  - 输出统一 utf-8
"""
import argparse
import io
import os
import re
import sys
import zipfile


def extract_docx_python_docx(path):
    import docx  # python-docx

    doc = docx.Document(path)
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_docx_zip(path):
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    xml = re.sub(r"</w:p>", "\n", xml)
    xml = re.sub(r"</w:tr>", "\n", xml)
    xml = re.sub(r"</w:tc>", " | ", xml)
    xml = re.sub(r"<[^>]+>", "", xml)
    return xml


def extract_docx(path):
    try:
        return extract_docx_python_docx(path)
    except Exception as e:  # noqa: BLE001
        print("[extract_text] python-docx 失败（%s），降级 zipfile 解析" % e, file=sys.stderr)
        return extract_docx_zip(path)


def extract_pdf(path):
    from pypdf import PdfReader

    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    text = "\n".join(parts).strip()
    if not text:
        print(
            "[extract_text] 错误：PDF 提取结果为空（可能是扫描件/图片型 PDF）。\n"
            "请将内容另存为 txt 后重试，或直接粘贴文本。",
            file=sys.stderr,
        )
        sys.exit(2)
    return text


def extract_txt(path):
    raw = open(path, "rb").read()
    for enc in ("utf-8", "gb18030"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def main():
    ap = argparse.ArgumentParser(description="提取 PDF/DOCX/TXT 为纯文本")
    ap.add_argument("--input", required=True)
    ap.add_argument("--output", required=True)
    args = ap.parse_args()

    ext = os.path.splitext(args.input)[1].lower()
    if ext == ".docx":
        text = extract_docx(args.input)
    elif ext == ".pdf":
        text = extract_pdf(args.input)
    elif ext in (".txt", ".md"):
        text = extract_txt(args.input)
    else:
        print("[extract_text] 错误：不支持的格式 %s（仅支持 pdf/docx/txt）" % ext, file=sys.stderr)
        sys.exit(2)

    text = text.strip()
    if not text:
        print("[extract_text] 错误：提取结果为空", file=sys.stderr)
        sys.exit(2)

    with io.open(args.output, "w", encoding="utf-8") as f:
        f.write(text + "\n")
    print("[extract_text] OK %s -> %s（%d 字符）" % (args.input, args.output, len(text)))


if __name__ == "__main__":
    main()
